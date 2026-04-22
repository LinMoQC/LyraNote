from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

from fix_thesis_er_diagrams import (
    ENTITY_SPECS,
    OUTPUT_DIR,
    REPO_ROOT,
    delete_paragraph,
    ensure_dirs,
    insert_paragraph_after,
    resolve_style,
)


SOURCE_DOC = OUTPUT_DIR / "220501020064-黄凯-毕业论文-引用修订版_4-ER图修订版.docx"
TMP_DIR = REPO_ROOT / "tmp" / "docs" / "thesis-er-diagram-mermaid"
OUTPUT_DOC = OUTPUT_DIR / "220501020064-黄凯-毕业论文-引用修订版_4-ER图Mermaid版.docx"

BG_COLOR = "#ffffff"
BOX_FILL = "#fff4dd"
BOX_BORDER = "#8b6f47"
HEADER_TEXT = "#3b342c"
ROW_ODD = "#ffffff"
ROW_EVEN = "#f2f2f2"
GRID = "#b9a88a"
TEXT = "#333333"
TYPE_TEXT = "#59534e"
KEY_FILL = "#ece8ff"
KEY_TEXT = "#51459b"


def pick_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        ("/System/Library/Fonts/PingFang.ttc", 1 if bold else 0),
        ("/System/Library/Fonts/STHeiti Medium.ttc", 0),
        ("/System/Library/Fonts/Supplemental/Songti.ttc", 0),
    ]
    for candidate, index in font_candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size, index=index)
            except OSError:
                continue
    return ImageFont.load_default()


def measure(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def parse_attr(attr: str) -> tuple[str, str, str]:
    key = ""
    name = attr
    if attr.endswith("(PK)"):
        name = attr[:-4]
        key = "PK"
    elif attr.endswith("(UK)"):
        name = attr[:-4]
        key = "UK"

    attr_type = infer_type(name)
    return attr_type, name, key


def infer_type(name: str) -> str:
    if name.endswith("ID"):
        return "uuid"
    if name.startswith("是否") or name in {"已读状态", "冲突标记"}:
        return "bool"
    if "时间" in name or "日期" in name:
        return "datetime"
    if "数量" in name or "次数" in name or "序号" in name or "版本号" in name or "耗时" in name:
        return "int"
    if "置信度" in name:
        return "float"
    if "嵌入" in name:
        return "vector"
    if (
        "JSON" in name
        or "元数据" in name
        or "列表" in name
        or "上下文" in name
        or name in {
            "参数配置",
            "投递配置",
            "核心主题",
            "活动上下文",
            "引用信息",
            "Agent步骤",
            "附件信息",
            "速度指标",
            "思维导图",
            "图表数据",
            "MCP结果",
            "UI元素",
            "画像JSON",
            "交付物JSON",
            "时间线JSON",
            "投递状态",
        }
    ):
        return "json"
    if (
        "内容" in name
        or "正文" in name
        or "摘要" in name
        or "描述" in name
        or "报告" in name
        or "错误" in name
        or "证据" in name
        or "轨迹" in name
        or "结果" in name
        or "文本" in name
    ):
        return "text"
    return "string"


def make_mermaid_source(entity_name: str, attrs: tuple[str, ...]) -> str:
    rows = []
    for attr in attrs:
        attr_type, name, key = parse_attr(attr)
        if key:
            rows.append(f"        {attr_type} {name} {key}")
        else:
            rows.append(f"        {attr_type} {name}")
    return "erDiagram\n" + f"    {entity_name} {{\n" + "\n".join(rows) + "\n    }\n"


def render_mermaid_style_entity(entity_name: str, attrs: tuple[str, ...]) -> tuple[Path, tuple[int, int]]:
    TMP_DIR.mkdir(parents=True, exist_ok=True)

    rows = [parse_attr(attr) for attr in attrs]
    mmd_path = TMP_DIR / f"{entity_name}.mmd"
    mmd_path.write_text(make_mermaid_source(entity_name, attrs), encoding="utf-8")

    title_font = pick_font(28, bold=True)
    row_font = pick_font(24)
    key_font = pick_font(18, bold=True)

    probe = Image.new("RGB", (10, 10), BG_COLOR)
    draw = ImageDraw.Draw(probe)

    type_width = max(measure(draw, r[0], row_font)[0] for r in rows)
    name_width = max(measure(draw, r[1], row_font)[0] for r in rows)
    header_width = measure(draw, entity_name, title_font)[0]
    has_key = any(r[2] for r in rows)

    cell_pad_x = 18
    header_h = 52
    row_h = 40
    border = 3
    key_w = 52 if has_key else 0
    type_col_w = type_width + cell_pad_x * 2
    name_col_w = max(name_width + cell_pad_x * 2, header_width + 48)
    table_w = type_col_w + name_col_w + key_w
    table_h = header_h + len(rows) * row_h
    image_w = table_w + border * 2
    image_h = table_h + border * 2

    image = Image.new("RGB", (image_w, image_h), BG_COLOR)
    draw = ImageDraw.Draw(image)

    left = border
    top = border
    right = left + table_w
    bottom = top + table_h
    draw.rounded_rectangle(
        (left, top, right, bottom),
        radius=10,
        fill=BOX_FILL,
        outline=BOX_BORDER,
        width=border,
    )
    draw.rectangle((left, top, right, top + header_h), fill=BOX_FILL, outline=None)
    draw.line((left, top + header_h, right, top + header_h), fill=BOX_BORDER, width=2)

    header_text_w, header_text_h = measure(draw, entity_name, title_font)
    draw.text(
        (left + (table_w - header_text_w) / 2, top + (header_h - header_text_h) / 2 - 2),
        entity_name,
        fill=HEADER_TEXT,
        font=title_font,
    )

    for idx, (attr_type, name, key) in enumerate(rows):
        row_top = top + header_h + idx * row_h
        row_bottom = row_top + row_h
        draw.rectangle(
            (left, row_top, right, row_bottom),
            fill=ROW_EVEN if idx % 2 else ROW_ODD,
            outline=None,
        )
        if idx < len(rows) - 1:
            draw.line((left, row_bottom, right, row_bottom), fill=GRID, width=1)

        draw.text((left + cell_pad_x, row_top + 7), attr_type, fill=TYPE_TEXT, font=row_font)
        draw.text((left + type_col_w + cell_pad_x, row_top + 7), name, fill=TEXT, font=row_font)

        draw.line((left + type_col_w, row_top, left + type_col_w, row_bottom), fill=GRID, width=1)
        if has_key:
            draw.line((right - key_w, row_top, right - key_w, row_bottom), fill=GRID, width=1)
            if key:
                pill_w = 34
                pill_h = 22
                pill_x = right - key_w + (key_w - pill_w) / 2
                pill_y = row_top + (row_h - pill_h) / 2
                draw.rounded_rectangle(
                    (pill_x, pill_y, pill_x + pill_w, pill_y + pill_h),
                    radius=10,
                    fill=KEY_FILL,
                    outline=None,
                )
                key_text_w, key_text_h = measure(draw, key, key_font)
                draw.text(
                    (pill_x + (pill_w - key_text_w) / 2, pill_y + (pill_h - key_text_h) / 2 - 1),
                    key,
                    fill=KEY_TEXT,
                    font=key_font,
                )

    png_path = TMP_DIR / f"{entity_name}.png"
    image.save(png_path)
    return png_path, (image_w, image_h)


def rebuild_er_section(document: Document, image_meta: dict[str, tuple[Path, tuple[int, int]]]) -> None:
    body_style = resolve_style(document, "毕业设计（论文）正文")
    caption_style = resolve_style(document, "图标", "Normal")
    image_style = resolve_style(document, "No Spacing", "Normal")

    document.paragraphs[183].text = (
        "经过对系统分析并结合当前代码中的数据模型定义，本文将用户、笔记本、知识来源、知识分块、对话、消息、"
        "深度研究任务、定时任务、用户记忆、用户画像、主动洞察、系统配置、笔记本摘要、会话摘要、定时任务运行记录和后台思考"
        "共 16 个核心实体整理为属性型 E-R 图。为突出实体自身语义，图中只保留实体固有属性，不再将外键字段纳入属性列表；"
        "本次图形表现采用 Mermaid 风格的实体卡片形式，以便在版面上更紧凑、更统一地呈现各实体结构。"
    )
    document.paragraphs[183].style = body_style

    document.paragraphs[189].text = (
        "在实体-关系（E-R）模型设计中，本文主要围绕关系型数据库 PostgreSQL 的核心业务实体进行建模。"
        "各实体的自有属性直接决定数据库表中的主要字段设计，而实体之间的关系则由主键、外键和关联约束进行表达。"
        "因此，以下实体图仅展示主键与非外键业务属性，不再把归属关系字段重复放入属性图中；"
        "同时采用 Mermaid 风格的表格式实体表示，以减少图片冗余留白并提升可读性。"
    )
    document.paragraphs[189].style = body_style

    for paragraph in list(document.paragraphs[190:238]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[189]
    for spec in ENTITY_SPECS:
        desc_para = insert_paragraph_after(anchor, spec.description, body_style)
        anchor = desc_para

        image_para = insert_paragraph_after(anchor, style=image_style)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_path, (px_w, _px_h) = image_meta[spec.figure_no]
        width_inches = min(5.8, max(2.6, px_w / 230))
        image_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
        anchor = image_para

        caption_para = insert_paragraph_after(anchor, spec.caption, caption_style)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = caption_para


def validate_output(document_path: Path) -> None:
    document = Document(document_path)
    expected = {spec.caption for spec in ENTITY_SPECS}
    actual = {p.text.strip() for p in document.paragraphs if p.text.strip().startswith("图3-")}
    missing = expected - actual
    if missing:
        raise RuntimeError(f"输出文档缺少图题: {sorted(missing)}")


def main() -> None:
    ensure_dirs()
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"未找到输入文档: {SOURCE_DOC}")

    image_meta = {}
    for spec in ENTITY_SPECS:
        image_meta[spec.figure_no] = render_mermaid_style_entity(spec.entity_name, spec.attrs)

    document = Document(SOURCE_DOC)
    rebuild_er_section(document, image_meta)
    document.save(OUTPUT_DOC)
    validate_output(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
