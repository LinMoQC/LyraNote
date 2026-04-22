from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


HEADERS = ["字段名称", "类型", "长度", "字段说明"]
COLUMN_WIDTHS_CM = (4.2, 2.1, 1.6, 7.1)
PRIMARY_OUTPUT_DOCX = Path("output/doc/LyraNote数据库表设计.docx")
VERSIONED_OUTPUT_DOCX = Path("output/doc/LyraNote数据库表设计-三线表-最新模型.docx")

TABLES = [
    (
        "表3-1 用户信息表（users）",
        [
            ("id", "UUID", "36", "用户主键"),
            ("username", "varchar", "255", "用户名"),
            ("password_hash", "varchar", "255", "密码哈希"),
            ("email", "varchar", "255", "邮箱"),
            ("name", "varchar", "255", "用户姓名/昵称"),
            ("avatar_url", "text", "-", "用户头像地址"),
            ("google_id", "varchar", "255", "Google 登录标识"),
            ("github_id", "varchar", "255", "GitHub 登录标识"),
            ("oauth_unbound", "varchar", "64", "已解绑第三方登录标记"),
            ("created_at", "datetime", "-", "创建时间"),
        ],
    ),
    (
        "表3-2 应用配置表（app_config）",
        [
            ("key", "varchar", "255", "配置项主键"),
            ("value", "text", "-", "配置项值"),
        ],
    ),
    (
        "表3-3 笔记本信息表（notebooks）",
        [
            ("id", "UUID", "36", "笔记本主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("title", "varchar", "500", "笔记本标题"),
            ("description", "text", "-", "笔记本描述"),
            ("status", "varchar", "50", "笔记本状态"),
            ("is_global", "boolean", "-", "是否全局知识库"),
            ("is_system", "boolean", "-", "是否系统笔记本"),
            ("system_type", "varchar", "50", "系统笔记本类型"),
            ("source_count", "integer", "-", "来源数量"),
            ("is_public", "boolean", "-", "是否公开"),
            ("published_at", "datetime", "-", "发布时间"),
            ("cover_emoji", "varchar", "10", "封面表情"),
            ("cover_gradient", "varchar", "50", "封面渐变样式"),
            ("created_at", "datetime", "-", "创建时间"),
            ("updated_at", "datetime", "-", "更新时间"),
        ],
    ),
    (
        "表3-4 笔记本摘要表（notebook_summaries）",
        [
            ("notebook_id", "UUID", "36", "所属笔记本 ID"),
            ("summary_md", "text", "-", "Markdown 摘要内容"),
            ("key_themes", "json", "-", "关键主题列表"),
            ("last_synced_at", "datetime", "-", "最近同步时间"),
        ],
    ),
    (
        "表3-5 知识来源表（sources）",
        [
            ("id", "UUID", "36", "来源主键"),
            ("notebook_id", "UUID", "36", "所属笔记本 ID"),
            ("title", "varchar", "500", "来源标题"),
            ("type", "varchar", "50", "来源类型"),
            ("status", "varchar", "50", "索引状态"),
            ("file_path", "text", "-", "本地文件路径"),
            ("url", "text", "-", "网页地址"),
            ("raw_text", "text", "-", "抽取后的原始文本"),
            ("summary", "text", "-", "来源摘要"),
            ("storage_key", "varchar", "500", "对象存储键"),
            ("storage_backend", "varchar", "20", "存储后端类型"),
            ("created_at", "datetime", "-", "创建时间"),
            ("updated_at", "datetime", "-", "更新时间"),
        ],
    ),
    (
        "表3-6 知识分块表（chunks）",
        [
            ("id", "UUID", "36", "分块主键"),
            ("source_id", "UUID", "36", "所属来源 ID"),
            ("notebook_id", "UUID", "36", "所属笔记本 ID"),
            ("content", "text", "-", "分块正文"),
            ("chunk_index", "integer", "-", "分块顺序号"),
            ("embedding", "vector", "-", "向量表示"),
            ("token_count", "integer", "-", "Token 数量"),
            ("metadata", "json", "-", "分块元数据"),
            ("created_at", "datetime", "-", "创建时间"),
            ("source_type", "varchar", "20", "分块来源类型"),
            ("note_id", "UUID", "36", "关联笔记 ID"),
        ],
    ),
    (
        "表3-7 对话信息表（conversations）",
        [
            ("id", "UUID", "36", "对话主键"),
            ("notebook_id", "UUID", "36", "关联笔记本 ID"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("title", "varchar", "500", "对话标题"),
            ("source", "varchar", "20", "对话来源"),
            ("created_at", "datetime", "-", "创建时间"),
        ],
    ),
    (
        "表3-8 消息信息表（messages）",
        [
            ("id", "UUID", "36", "消息主键"),
            ("conversation_id", "UUID", "36", "所属对话 ID"),
            ("generation_id", "UUID", "36", "关联生成任务 ID"),
            ("role", "varchar", "50", "消息角色"),
            ("status", "varchar", "20", "消息状态"),
            ("content", "text", "-", "消息正文"),
            ("reasoning", "text", "-", "推理内容"),
            ("citations", "json", "-", "引用信息"),
            ("agent_steps", "json", "-", "Agent 执行步骤"),
            ("attachments", "json", "-", "附件信息"),
            ("speed", "json", "-", "流式速度指标"),
            ("mind_map", "json", "-", "思维导图结果"),
            ("diagram", "json", "-", "图形结果"),
            ("mcp_result", "json", "-", "MCP 工具结果"),
            ("ui_elements", "json", "-", "生成式界面元素"),
            ("parent_message_id", "UUID", "36", "父消息 ID"),
            ("created_at", "datetime", "-", "创建时间"),
        ],
    ),
    (
        "表3-9 对话摘要表（conversation_summaries）",
        [
            ("conversation_id", "UUID", "36", "所属对话 ID"),
            ("summary_text", "text", "-", "压缩后的对话摘要"),
            ("compressed_message_count", "integer", "-", "已压缩消息数量"),
            ("compressed_through", "datetime", "-", "摘要覆盖截止时间"),
            ("updated_at", "datetime", "-", "更新时间"),
        ],
    ),
    (
        "表3-10 深度研究任务表（research_tasks）",
        [
            ("id", "UUID", "36", "任务主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("notebook_id", "varchar", "100", "关联笔记本 ID"),
            ("conversation_id", "UUID", "36", "关联对话 ID"),
            ("query", "text", "-", "研究问题"),
            ("mode", "varchar", "20", "研究模式"),
            ("status", "varchar", "20", "任务状态"),
            ("report", "text", "-", "研究报告"),
            ("deliverable_json", "json", "-", "结构化交付物"),
            ("timeline_json", "json", "-", "过程时间线"),
            ("web_sources_json", "json", "-", "外部资料列表"),
            ("error_message", "text", "-", "错误信息"),
            ("created_at", "datetime", "-", "创建时间"),
            ("completed_at", "datetime", "-", "完成时间"),
        ],
    ),
    (
        "表3-11 定时任务表（scheduled_tasks）",
        [
            ("id", "UUID", "36", "定时任务主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("name", "varchar", "255", "任务名称"),
            ("description", "text", "-", "任务描述"),
            ("task_type", "varchar", "50", "任务类型"),
            ("schedule_cron", "varchar", "100", "调度表达式"),
            ("timezone", "varchar", "50", "时区"),
            ("parameters", "json", "-", "任务参数"),
            ("delivery_config", "json", "-", "投递配置"),
            ("enabled", "boolean", "-", "是否启用"),
            ("last_run_at", "datetime", "-", "上次执行时间"),
            ("next_run_at", "datetime", "-", "下次执行时间"),
            ("run_count", "integer", "-", "执行次数"),
            ("last_result", "text", "-", "最近执行结果"),
            ("last_error", "text", "-", "最近错误信息"),
            ("consecutive_failures", "integer", "-", "连续失败次数"),
            ("created_at", "datetime", "-", "创建时间"),
            ("updated_at", "datetime", "-", "更新时间"),
        ],
    ),
    (
        "表3-12 定时任务执行记录表（scheduled_task_runs）",
        [
            ("id", "UUID", "36", "执行记录主键"),
            ("task_id", "UUID", "36", "所属定时任务 ID"),
            ("status", "varchar", "20", "执行状态"),
            ("started_at", "datetime", "-", "开始时间"),
            ("finished_at", "datetime", "-", "结束时间"),
            ("duration_ms", "integer", "-", "执行耗时（毫秒）"),
            ("result_summary", "text", "-", "结果摘要"),
            ("error_message", "text", "-", "错误信息"),
            ("generated_content", "text", "-", "生成内容"),
            ("sources_count", "integer", "-", "涉及来源数量"),
            ("delivery_status", "json", "-", "投递状态"),
        ],
    ),
    (
        "表3-13 用户记忆表（user_memories）",
        [
            ("id", "UUID", "36", "用户记忆主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("key", "varchar", "100", "记忆键"),
            ("value", "text", "-", "记忆值"),
            ("confidence", "float", "-", "置信度"),
            ("memory_type", "varchar", "20", "记忆类型"),
            ("memory_kind", "varchar", "32", "记忆类别"),
            ("access_count", "integer", "-", "访问次数"),
            ("last_accessed_at", "datetime", "-", "最后访问时间"),
            ("expires_at", "datetime", "-", "过期时间"),
            ("reinforced_by", "varchar", "36", "强化来源反思 ID"),
            ("updated_at", "datetime", "-", "更新时间"),
            ("embedding", "vector", "-", "记忆向量"),
            ("source", "varchar", "20", "记忆来源类型"),
            ("evidence", "text", "-", "证据链或来源记录"),
            ("conflict_flag", "boolean", "-", "冲突标记"),
        ],
    ),
    (
        "表3-14 用户画像表（user_portraits）",
        [
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("portrait_json", "json", "-", "六维画像 JSON"),
            ("avatar_url", "text", "-", "画像头像地址"),
            ("synthesis_summary", "text", "-", "画像合成摘要"),
            ("version", "integer", "-", "画像版本号"),
            ("synthesized_at", "datetime", "-", "合成时间"),
            ("created_at", "datetime", "-", "创建时间"),
            ("updated_at", "datetime", "-", "更新时间"),
        ],
    ),
    (
        "表3-15 Agent 思想记录表（agent_thoughts）",
        [
            ("id", "UUID", "36", "思想记录主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("visibility", "varchar", "20", "可见性状态"),
            ("content", "text", "-", "思想内容"),
            ("activity_context", "json", "-", "活动上下文快照"),
            ("notebook_id", "UUID", "36", "关联笔记本 ID"),
            ("created_at", "datetime", "-", "创建时间"),
        ],
    ),
    (
        "表3-16 主动洞察表（proactive_insights）",
        [
            ("id", "UUID", "36", "洞察主键"),
            ("user_id", "UUID", "36", "所属用户 ID"),
            ("notebook_id", "UUID", "36", "关联笔记本 ID"),
            ("insight_type", "varchar", "50", "洞察类型"),
            ("title", "varchar", "500", "洞察标题"),
            ("content", "text", "-", "洞察内容"),
            ("metadata", "json", "-", "洞察附加元数据"),
            ("is_read", "boolean", "-", "是否已读"),
            ("created_at", "datetime", "-", "创建时间"),
        ],
    ),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_doc_fonts(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_cell_margins(cell, top: int = 40, start: int = 60, bottom: int = 40, end: int = 60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, **borders) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge, data in borders.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    styles = {
        "top": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "bottom": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "left": {"val": "nil"},
        "right": {"val": "nil"},
        "insideH": {"val": "nil"},
        "insideV": {"val": "nil"},
    }
    for edge, data in styles.items():
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        for key, value in data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def format_three_line_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for row in table.rows:
        set_row_cant_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.8)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, COLUMN_WIDTHS_CM[index])
            set_cell_margins(cell)
            clear_cell_borders(cell)

    header_border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom=header_border)


def add_caption(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    set_east_asia_font(run, "宋体")


def add_intro_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    set_east_asia_font(run, "宋体")


def add_table(doc: Document, title: str, rows: list[tuple[str, str, str, str]]) -> None:
    add_caption(doc, title)

    table = doc.add_table(rows=len(rows) + 1, cols=len(HEADERS))
    for index, header in enumerate(HEADERS):
        set_cell_text(table.cell(0, index), header, bold=True)

    for row_index, row_data in enumerate(rows, start=1):
        for col_index, value in enumerate(row_data):
            set_cell_text(table.cell(row_index, col_index), value)

    format_three_line_table(table)
    doc.add_paragraph("")


def build_doc() -> Document:
    doc = Document()
    set_doc_fonts(doc)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("3.3.2 数据库表设计")
    title_run.bold = True
    title_run.font.size = Pt(14)
    set_east_asia_font(title_run, "黑体")

    add_intro_paragraph(
        doc,
        "根据 LyraNote 系统当前代码实现，数据库表主要围绕用户管理、知识组织、对话交互、研究任务、长期记忆和主动服务等几个方面展开设计。系统底层采用 PostgreSQL 作为核心关系型数据库，并结合 pgvector 扩展保存知识分块与长期记忆的向量表示。",
    )
    add_intro_paragraph(
        doc,
        "为体现当前版本的数据结构设计，本文选取用户、应用配置、笔记本、笔记本摘要、知识来源、知识分块、对话、消息、对话摘要、深度研究任务、定时任务、定时任务执行记录、用户记忆、用户画像、Agent 思想记录和主动洞察等 16 张核心数据表进行说明。各主要数据表设计如下所示。",
    )

    for title_text, rows in TABLES:
        add_table(doc, title_text, rows)

    return doc


def main() -> None:
    PRIMARY_OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(PRIMARY_OUTPUT_DOCX)
    doc.save(VERSIONED_OUTPUT_DOCX)
    print(PRIMARY_OUTPUT_DOCX)
    print(VERSIONED_OUTPUT_DOCX)


if __name__ == "__main__":
    main()
