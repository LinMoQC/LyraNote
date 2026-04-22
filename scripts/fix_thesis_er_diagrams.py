from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = Path("/Users/kaihuang/Documents/毕业论文/220501020064-黄凯-毕业论文-引用修订版_4.docx")
TMP_DIR = REPO_ROOT / "tmp" / "docs" / "thesis-er-diagram-cleanup"
OUTPUT_DIR = REPO_ROOT / "output" / "doc"
OUTPUT_DOC = OUTPUT_DIR / "220501020064-黄凯-毕业论文-引用修订版_4-ER图修订版.docx"


@dataclass(frozen=True)
class EntitySpec:
    index: int
    figure_no: str
    entity_name: str
    caption: str
    description: str
    attrs: tuple[str, ...]

    @property
    def image_name(self) -> str:
        return f"{self.figure_no}-{self.entity_name}.png"


ENTITY_SPECS: tuple[EntitySpec, ...] = (
    EntitySpec(
        index=1,
        figure_no="3-6",
        entity_name="用户",
        caption="图3-6 用户实体 E-R 图",
        description=(
            "（1）用户实体 E-R 图如图3-6所示。图 3-6 展示了用户实体的 E-R 结构。"
            "用户实体用于标识系统中的真实使用者，图中仅保留用户自身属性，"
            "包括账号信息、认证标识、头像信息以及创建时间等内容；"
            "与笔记本、会话、任务和长期记忆等实体之间的归属关系在后续关系约束中体现。"
        ),
        attrs=(
            "用户ID(PK)",
            "用户名",
            "密码哈希",
            "邮箱",
            "姓名",
            "头像地址",
            "Google标识",
            "GitHub标识",
            "OAuth解绑标记",
            "创建时间",
        ),
    ),
    EntitySpec(
        index=2,
        figure_no="3-7",
        entity_name="笔记本",
        caption="图3-7 笔记本实体 E-R 图",
        description=(
            "（2）笔记本实体 E-R 图如图3-7所示。图 3-7 展示了笔记本实体的 E-R 结构。"
            "笔记本是 LyraNote 组织私有知识的核心工作空间，"
            "图中保留了标题、描述、状态、公开性与封面等实体自有属性，"
            "而与用户、来源、对话和笔记等对象的关联不再作为属性列重复展示。"
        ),
        attrs=(
            "笔记本ID(PK)",
            "标题",
            "描述",
            "状态",
            "是否全局知识库",
            "是否系统笔记本",
            "系统类型",
            "来源数量",
            "是否公开",
            "发布时间",
            "封面表情",
            "封面渐变",
            "创建时间",
            "更新时间",
        ),
    ),
    EntitySpec(
        index=3,
        figure_no="3-8",
        entity_name="知识来源",
        caption="图3-8 知识来源实体 E-R 图",
        description=(
            "（3）知识来源实体 E-R 图如图3-8所示。图 3-8 展示了知识来源实体的 E-R 结构。"
            "知识来源记录用户导入的 PDF、网页和 Markdown 等资料，"
            "实体自身属性主要描述来源类型、索引状态、文本内容、摘要和存储信息，"
            "其与笔记本及知识分块之间的联系通过关系设计表达。"
        ),
        attrs=(
            "来源ID(PK)",
            "标题",
            "来源类型",
            "索引状态",
            "文件路径",
            "链接地址",
            "原始文本",
            "摘要",
            "存储键",
            "存储后端",
            "创建时间",
            "更新时间",
        ),
    ),
    EntitySpec(
        index=4,
        figure_no="3-9",
        entity_name="知识分块",
        caption="图3-9 知识分块实体 E-R 图",
        description=(
            "（4）知识分块实体 E-R 图如图3-9所示。图 3-9 展示了知识分块实体的 E-R 结构。"
            "知识分块是系统执行检索增强生成时使用的最小知识单元，"
            "图中仅保留分块正文、顺序编号、向量嵌入、元数据和来源类型等自身属性，"
            "不再将所属来源、笔记本或笔记的外键字段混入实体属性。"
        ),
        attrs=(
            "分块ID(PK)",
            "正文内容",
            "分块序号",
            "向量嵌入",
            "Token数量",
            "元数据",
            "来源类型",
            "创建时间",
        ),
    ),
    EntitySpec(
        index=5,
        figure_no="3-10",
        entity_name="对话",
        caption="图3-10 对话实体 E-R 图",
        description=(
            "（5）对话实体 E-R 图如图3-10所示。图 3-10 展示了对话实体的 E-R 结构。"
            "对话实体用于承载一次连续的交互会话，"
            "其自身属性主要包括对话标题、来源场景和创建时间；"
            "与用户、笔记本以及消息记录的关联关系通过实体联系表达，而不直接写入属性列表。"
        ),
        attrs=(
            "对话ID(PK)",
            "标题",
            "对话来源",
            "创建时间",
        ),
    ),
    EntitySpec(
        index=6,
        figure_no="3-11",
        entity_name="消息",
        caption="图3-11 消息实体 E-R 图",
        description=(
            "（6）消息实体 E-R 图如图3-11所示。图 3-11 展示了消息实体的 E-R 结构。"
            "消息实体是对话过程中的最小交互单元，"
            "保存角色、状态、正文、引用、推理轨迹以及富媒体结果等内容；"
            "对话归属、生成归属和父消息分支等外键信息不再作为本实体属性展示。"
        ),
        attrs=(
            "消息ID(PK)",
            "角色",
            "状态",
            "正文内容",
            "推理轨迹",
            "引用信息",
            "Agent步骤",
            "附件信息",
            "速度指标",
            "思维导图",
            "图表数据",
            "MCP结果",
            "UI元素",
            "创建时间",
        ),
    ),
    EntitySpec(
        index=7,
        figure_no="3-12",
        entity_name="深度研究任务",
        caption="图3-12 深度研究任务实体 E-R 图",
        description=(
            "（7）深度研究任务实体 E-R 图如图3-12所示。图 3-12 展示了深度研究任务实体的 E-R 结构。"
            "深度研究任务用于表示一次长生命周期的研究流程，"
            "实体自身属性包括研究问题、执行模式、状态、报告正文、结构化交付物和时间线等，"
            "与用户、会话和笔记本的关联通过关系约束体现。"
        ),
        attrs=(
            "任务ID(PK)",
            "研究问题",
            "研究模式",
            "状态",
            "研究报告",
            "交付物JSON",
            "时间线JSON",
            "外部资料列表",
            "错误信息",
            "创建时间",
            "完成时间",
        ),
    ),
    EntitySpec(
        index=8,
        figure_no="3-13",
        entity_name="定时任务",
        caption="图3-13 定时任务实体 E-R 图",
        description=(
            "（8）定时任务实体 E-R 图如图3-13所示。图 3-13 展示了定时任务实体的 E-R 结构。"
            "定时任务实体描述用户配置的持续服务模板，"
            "图中保留任务名称、调度表达式、参数配置、投递配置和运行状态等自身属性，"
            "而用户归属关系仅通过实体联系表示，不再以外键字段形式放入属性图。"
        ),
        attrs=(
            "任务ID(PK)",
            "任务名称",
            "描述",
            "任务类型",
            "调度表达式",
            "时区",
            "参数配置",
            "投递配置",
            "是否启用",
            "最近执行时间",
            "下次执行时间",
            "执行次数",
            "最近结果",
            "最近错误",
            "连续失败次数",
            "创建时间",
            "更新时间",
        ),
    ),
    EntitySpec(
        index=9,
        figure_no="3-14",
        entity_name="用户记忆",
        caption="图3-14 用户记忆实体 E-R 图",
        description=(
            "（9）用户记忆实体 E-R 图如图3-14所示。图 3-14 展示了用户记忆实体的 E-R 结构。"
            "用户记忆用于保存系统长期积累的偏好、事实与技能画像，"
            "实体自身属性包括键值内容、置信度、访问统计、失效时间、向量嵌入和证据来源等，"
            "用户归属关系在关系层处理，不再混入属性集合。"
        ),
        attrs=(
            "记忆ID(PK)",
            "记忆键",
            "记忆值",
            "置信度",
            "记忆类型",
            "记忆类别",
            "访问次数",
            "最近访问时间",
            "失效时间",
            "强化来源",
            "更新时间",
            "向量嵌入",
            "来源类型",
            "证据",
            "冲突标记",
        ),
    ),
    EntitySpec(
        index=10,
        figure_no="3-15",
        entity_name="用户画像",
        caption="图3-15 用户画像实体 E-R 图",
        description=(
            "（10）用户画像实体 E-R 图如图3-15所示。图 3-15 展示了用户画像实体的 E-R 结构。"
            "用户画像实体是对长期记忆与交互行为进行阶段性聚合后的结果，"
            "图中仅保留画像内容、头像地址、合成摘要、版本号以及时间相关属性，"
            "一对一归属关系由实体联系表达而不作为属性列出现。"
        ),
        attrs=(
            "画像JSON",
            "头像地址",
            "合成摘要",
            "版本号",
            "合成时间",
            "创建时间",
            "更新时间",
        ),
    ),
    EntitySpec(
        index=11,
        figure_no="3-16",
        entity_name="主动洞察",
        caption="图3-16 主动洞察实体 E-R 图",
        description=(
            "（11）主动洞察实体 E-R 图如图3-16所示。图 3-16 展示了主动洞察实体的 E-R 结构。"
            "主动洞察实体用于保存可以直接展示给用户的系统提示内容，"
            "其自身属性包括洞察类型、标题、正文、元数据、已读状态与创建时间，"
            "与用户和笔记本的关联关系则由实体之间的联系体现。"
        ),
        attrs=(
            "洞察ID(PK)",
            "洞察类型",
            "标题",
            "内容",
            "元数据",
            "已读状态",
            "创建时间",
        ),
    ),
    EntitySpec(
        index=12,
        figure_no="3-17",
        entity_name="系统配置",
        caption="图3-17 系统配置实体 E-R 图",
        description=(
            "（12）系统配置实体 E-R 图如图3-17所示。图 3-17 展示了系统配置实体的 E-R 结构。"
            "系统配置实体采用键值形式保存安装向导和运行期写入的配置项，"
            "由于该实体本身不承担业务归属关系，因此图中仅包含配置键与配置值两个自身属性。"
        ),
        attrs=(
            "配置键(PK)",
            "配置值",
        ),
    ),
    EntitySpec(
        index=13,
        figure_no="3-18",
        entity_name="笔记本摘要",
        caption="图3-18 笔记本摘要实体 E-R 图",
        description=(
            "（13）笔记本摘要实体 E-R 图如图3-18所示。图 3-18 展示了笔记本摘要实体的 E-R 结构。"
            "笔记本摘要实体用于保存系统自动生成的笔记本级聚合摘要，"
            "图中仅展示摘要内容、主题词以及同步时间等实体自身属性，"
            "其与笔记本的一对一绑定关系不再作为属性列重复绘制。"
        ),
        attrs=(
            "摘要内容",
            "核心主题",
            "最后同步时间",
        ),
    ),
    EntitySpec(
        index=14,
        figure_no="3-19",
        entity_name="会话摘要",
        caption="图3-19 会话摘要实体 E-R 图",
        description=(
            "（14）会话摘要实体 E-R 图如图3-19所示。图 3-19 展示了会话摘要实体的 E-R 结构。"
            "会话摘要实体用于压缩长对话中的历史消息，"
            "实体自身属性包含摘要正文、已压缩消息数量、压缩边界时间和更新时间等内容，"
            "与对话实体的一对一归属关系通过关系设计表示。"
        ),
        attrs=(
            "摘要正文",
            "压缩消息数",
            "压缩边界时间",
            "更新时间",
        ),
    ),
    EntitySpec(
        index=15,
        figure_no="3-20",
        entity_name="定时任务运行记录",
        caption="图3-20 定时任务运行记录实体 E-R 图",
        description=(
            "（15）定时任务运行记录实体 E-R 图如图3-20所示。图 3-20 展示了定时任务运行记录实体的 E-R 结构。"
            "该实体用于记录定时任务每一次实际执行的过程和结果，"
            "图中保留运行状态、开始与结束时间、执行耗时、结果摘要、错误信息以及投递状态等自身属性，"
            "所属任务关系不再写入属性图。"
        ),
        attrs=(
            "运行ID(PK)",
            "状态",
            "开始时间",
            "结束时间",
            "耗时毫秒",
            "结果摘要",
            "错误信息",
            "生成内容",
            "来源数量",
            "投递状态",
        ),
    ),
    EntitySpec(
        index=16,
        figure_no="3-21",
        entity_name="后台思考",
        caption="图3-21 后台思考实体 E-R 图",
        description=(
            "（16）后台思考实体 E-R 图如图3-21所示。图 3-21 展示了后台思考实体的 E-R 结构。"
            "后台思考实体用于保存 Lyra Soul 在后台循环中形成的思考记录，"
            "其自身属性包括可见性、思考内容、活动上下文和创建时间；"
            "与用户和笔记本的关联关系通过实体联系表达，不再作为属性字段直接绘制。"
        ),
        attrs=(
            "思考ID(PK)",
            "可见性",
            "思考内容",
            "活动上下文",
            "创建时间",
        ),
    ),
)


def ensure_dirs() -> None:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def pick_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for candidate in font_candidates:
        path = Path(candidate)
        if path.exists():
            try:
                index = 1 if bold and "PingFang" in path.name else 0
                return ImageFont.truetype(str(path), size=size, index=index)
            except OSError:
                continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> str:
    if draw.textbbox((0, 0), text, font=font)[2] <= max_width:
        return text

    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return "\n".join(lines[:2])


def chunk_counts(n: int) -> tuple[int, ...]:
    if n <= 6:
        return (n,)
    if n <= 12:
        first = (n + 1) // 2
        return (first, n - first)
    first = min(6, (n + 2) // 3)
    second = min(6, (n - first + 1) // 2)
    third = n - first - second
    return tuple(count for count in (first, second, third) if count)


def generate_diagram(spec: EntitySpec) -> Path:
    canvas_w = 2200
    canvas_h = 1300
    image = Image.new("RGB", (canvas_w, canvas_h), "#ffffff")
    draw = ImageDraw.Draw(image)

    title_font = pick_font(54, bold=True)
    box_font = pick_font(56, bold=True)
    attr_font = pick_font(34)

    box_w = 460
    box_h = 120
    counts = chunk_counts(len(spec.attrs))
    row_ys = [110, 300, 490][: len(counts)]
    box_top = 770 if len(counts) >= 3 else 710 if len(counts) == 2 else 650
    box_left = (canvas_w - box_w) // 2
    box_right = box_left + box_w
    box_bottom = box_top + box_h

    ellipse_w = 260
    ellipse_h = 92
    gap_x = 36

    draw.text((90, 40), f"图 {spec.figure_no}  {spec.entity_name} 实体", fill="#333333", font=title_font)

    draw.rounded_rectangle(
        (box_left, box_top, box_right, box_bottom),
        radius=18,
        outline="#8a7ae6",
        width=5,
        fill="#f7f2ff",
    )
    box_text = spec.entity_name
    text_bbox = draw.multiline_textbbox((0, 0), box_text, font=box_font, spacing=8, align="center")
    box_text_w = text_bbox[2] - text_bbox[0]
    box_text_h = text_bbox[3] - text_bbox[1]
    draw.multiline_text(
        ((canvas_w - box_text_w) / 2, box_top + (box_h - box_text_h) / 2 - 4),
        box_text,
        fill="#473a8b",
        font=box_font,
        align="center",
        spacing=8,
    )

    attr_iter = iter(spec.attrs)
    for row_count, row_y in zip(counts, row_ys):
        total_w = row_count * ellipse_w + (row_count - 1) * gap_x
        start_x = (canvas_w - total_w) // 2
        for i in range(row_count):
            attr = next(attr_iter)
            left = start_x + i * (ellipse_w + gap_x)
            top = row_y
            right = left + ellipse_w
            bottom = top + ellipse_h
            center_x = (left + right) // 2
            center_y = (top + bottom) // 2

            draw.line(
                [(center_x, bottom), (canvas_w // 2, box_top)],
                fill="#c6b9f8",
                width=3,
            )
            draw.ellipse(
                (left, top, right, bottom),
                outline="#9d8ff0",
                width=4,
                fill="#f3efff",
            )

            label = wrap_text(draw, attr, attr_font, ellipse_w - 24)
            bbox = draw.multiline_textbbox((0, 0), label, font=attr_font, spacing=4, align="center")
            text_w = bbox[2] - bbox[0]
            text_h = bbox[3] - bbox[1]
            draw.multiline_text(
                (center_x - text_w / 2, center_y - text_h / 2 - 2),
                label,
                fill="#2f2f2f",
                font=attr_font,
                align="center",
                spacing=4,
            )

    output = TMP_DIR / spec.image_name
    image.save(output)
    return output


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def resolve_style(document: Document, preferred: str, fallback: str = "Normal") -> str:
    styles = {style.name for style in document.styles}
    if preferred in styles:
        return preferred
    return fallback if fallback in styles else "Normal"


def rebuild_er_section(document: Document, image_paths: dict[str, Path]) -> None:
    body_style = resolve_style(document, "毕业设计（论文）正文")
    caption_style = resolve_style(document, "图标", "Normal")
    image_style = resolve_style(document, "No Spacing", "Normal")

    document.paragraphs[183].text = (
        "经过对系统分析并结合当前代码中的数据模型定义，"
        "本文将用户、笔记本、知识来源、知识分块、对话、消息、深度研究任务、定时任务、用户记忆、"
        "用户画像、主动洞察、系统配置、笔记本摘要、会话摘要、定时任务运行记录和后台思考共 16 个核心实体"
        "整理为属性型 E-R 图。为突出实体自身语义，图中只保留实体固有属性，"
        "不再将外键字段纳入属性列表，实体之间的归属和关联关系通过后续数据库约束说明体现。"
    )
    document.paragraphs[183].style = body_style

    document.paragraphs[189].text = (
        "在实体-关系（E-R）模型设计中，本文主要围绕关系型数据库 PostgreSQL 的核心业务实体进行建模。"
        "各实体的自有属性直接决定数据库表中的主要字段设计，而实体之间的关系则由主键、外键和关联约束进行表达。"
        "因此，以下实体图仅展示主键与非外键业务属性，"
        "不再把归属关系字段重复放入属性图中，以避免实体属性与关系属性混淆。"
    )
    document.paragraphs[189].style = body_style

    for paragraph in list(document.paragraphs[190:222]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[189]
    for spec in ENTITY_SPECS:
        desc_para = insert_paragraph_after(anchor, spec.description, body_style)
        anchor = desc_para

        image_para = insert_paragraph_after(anchor, style=image_style)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.add_run().add_picture(str(image_paths[spec.figure_no]), width=Inches(6.8))
        anchor = image_para

        caption_para = insert_paragraph_after(anchor, spec.caption, caption_style)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = caption_para


def validate_output(document_path: Path) -> None:
    document = Document(document_path)
    expected_captions = {spec.caption for spec in ENTITY_SPECS}
    actual_captions = {p.text.strip() for p in document.paragraphs if p.text.strip().startswith("图3-")}
    missing = expected_captions - actual_captions
    if missing:
        raise RuntimeError(f"输出文档缺少图题: {sorted(missing)}")


def main() -> None:
    ensure_dirs()
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"未找到输入文档: {SOURCE_DOC}")

    image_paths = {spec.figure_no: generate_diagram(spec) for spec in ENTITY_SPECS}

    document = Document(SOURCE_DOC)
    rebuild_er_section(document, image_paths)
    document.save(OUTPUT_DOC)
    validate_output(OUTPUT_DOC)

    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
